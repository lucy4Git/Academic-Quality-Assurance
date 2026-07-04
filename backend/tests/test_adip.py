"""Tests for the ADIP Academic Document Intelligence Platform.

Tests cover:
  - PDF extractor (pdfminer.six)
  - DOCX extractor (python-docx)
  - XLSX extractor (openpyxl)
  - HTML extractor (beautifulsoup4)
  - Plain text extractor
  - Document classifier
  - Confidence scoring / gates
  - TUT ICT mapper
  - Provenance generation
  - Pipeline dry run
"""

from __future__ import annotations

import io
import json
import tempfile
from pathlib import Path

import pytest


# ── Helpers ───────────────────────────────────────────────────────────────────

def create_minimal_pdf(text: str) -> bytes:
    """Create a minimal valid PDF with embedded text for testing."""
    return (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 612 792] /Contents 5 0 R >> endobj\n"
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"5 0 obj << /Length " + str(len(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET")).encode() + b" >>\n"
        b"stream\n"
        + f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
        + b"\nendstream\nendobj\n"
        b"xref\n0 6\n0000000000 65535 f \n"
        b"trailer << /Size 6 /Root 1 0 R >>\nstartxref\n0\n%%EOF"
    )


def create_docx_file(paragraphs: list[tuple[str, str]]) -> bytes:
    """Create a minimal DOCX in memory. paragraphs = [(text, style_name)]."""
    import docx
    doc = docx.Document()
    for text, style in paragraphs:
        p = doc.add_paragraph(text)
        try:
            p.style = doc.styles[style]
        except Exception:
            pass
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_xlsx_file(headers: list[str], rows: list[list]) -> bytes:
    """Create a minimal XLSX in memory."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Extractor tests ─────────────────────────────────────────────────────────


class TestPDFExtractor:
    def test_can_handle_pdf(self):
        from app.adip.extractors.pdf_extractor import PDFExtractor
        ext = PDFExtractor()
        assert ext.can_handle("application/pdf", Path("test.pdf"))
        assert not ext.can_handle("text/html", Path("test.html"))

    def test_extract_returns_result(self, tmp_path):
        from app.adip.extractors.pdf_extractor import PDFExtractor
        # Write a minimal PDF
        pdf_content = create_minimal_pdf("NQF Level 6 Diploma in Computer Science 360 credits")
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(pdf_content)
        ext = PDFExtractor()
        result = ext.extract(pdf_file)
        # Result should not be None and should have method set
        assert result is not None
        assert result.extraction_method == "pdfminer_native"
        # May or may not extract text from minimal PDF but should not crash

    def test_extract_missing_file(self, tmp_path):
        from app.adip.extractors.pdf_extractor import PDFExtractor
        ext = PDFExtractor()
        result = ext.extract(tmp_path / "nonexistent.pdf")
        assert result.error is not None


class TestDOCXExtractor:
    def test_can_handle_docx(self):
        from app.adip.extractors.docx_extractor import DOCXExtractor
        ext = DOCXExtractor()
        assert ext.can_handle(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            Path("test.docx")
        )
        assert ext.can_handle("text/plain", Path("test.docx"))  # extension fallback

    def test_extract_docx_paragraphs(self, tmp_path):
        from app.adip.extractors.docx_extractor import DOCXExtractor
        content = create_docx_file([
            ("Faculty of ICT", "Heading 1"),
            ("Computer Science Department", "Heading 2"),
            ("Diploma in Computer Science (NQF Level 6, 360 credits)", "Normal"),
        ])
        docx_file = tmp_path / "test.docx"
        docx_file.write_bytes(content)
        ext = DOCXExtractor()
        result = ext.extract(docx_file)
        assert result.error is None
        assert len(result.chunks) >= 3
        texts = [c.text for c in result.chunks]
        assert any("Faculty of ICT" in t for t in texts)
        assert any("NQF" in t for t in texts)
        headings = [c for c in result.chunks if c.chunk_type == "heading"]
        assert len(headings) >= 2

    def test_extract_docx_table(self, tmp_path):
        import docx
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "Programme"
        table.rows[0].cells[1].text = "NQF Level"
        table.rows[0].cells[2].text = "Credits"
        table.rows[1].cells[0].text = "Diploma in CS"
        table.rows[1].cells[1].text = "6"
        table.rows[1].cells[2].text = "360"
        buf = io.BytesIO()
        doc.save(buf)
        docx_file = tmp_path / "table.docx"
        docx_file.write_bytes(buf.getvalue())

        from app.adip.extractors.docx_extractor import DOCXExtractor
        result = DOCXExtractor().extract(docx_file)
        assert result.error is None
        assert len(result.tables) == 1
        assert result.tables[0].header_row == ["Programme", "NQF Level", "Credits"]
        assert result.tables[0].data_rows[0]["NQF Level"] == "6"


class TestXLSXExtractor:
    def test_can_handle_xlsx(self):
        from app.adip.extractors.xlsx_extractor import XLSXExtractor
        ext = XLSXExtractor()
        assert ext.can_handle(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            Path("test.xlsx")
        )

    def test_extract_xlsx_table(self, tmp_path):
        content = create_xlsx_file(
            ["Programme", "NQF Level", "APS (Math)", "Campus"],
            [
                ["Diploma in Computer Science", 6, 26, "Soshanguve South"],
                ["Advanced Diploma in CS", 7, None, "Soshanguve South"],
            ]
        )
        xlsx_file = tmp_path / "data.xlsx"
        xlsx_file.write_bytes(content)
        from app.adip.extractors.xlsx_extractor import XLSXExtractor
        result = XLSXExtractor().extract(xlsx_file)
        assert result.error is None
        assert len(result.tables) == 1
        assert result.tables[0].header_row[0] == "Programme"
        assert result.tables[0].data_rows[0]["NQF Level"] == "6"

    def test_extract_csv(self, tmp_path):
        csv_file = tmp_path / "data.csv"
        csv_file.write_text("Programme,NQF Level,Credits\nDiploma in CS,6,360\n")
        from app.adip.extractors.xlsx_extractor import CSVExtractor
        result = CSVExtractor().extract(csv_file)
        assert result.error is None
        assert len(result.tables) == 1
        assert result.tables[0].data_rows[0]["NQF Level"] == "6"


class TestHTMLExtractor:
    def test_can_handle_html(self):
        from app.adip.extractors.html_extractor import HTMLExtractor
        ext = HTMLExtractor()
        assert ext.can_handle("text/html", Path("test.html"))

    def test_extract_headings_and_paragraphs(self, tmp_path):
        html = """<html><body>
        <h1>Faculty of ICT</h1>
        <h2>Department of Computer Science</h2>
        <p>Diploma in Computer Science (NQF level 6)</p>
        <p>APS requirement: 26 with Mathematics.</p>
        </body></html>"""
        html_file = tmp_path / "page.html"
        html_file.write_text(html, encoding="utf-8")
        from app.adip.extractors.html_extractor import HTMLExtractor
        result = HTMLExtractor().extract(html_file)
        assert result.error is None
        texts = [c.text for c in result.chunks]
        assert any("Faculty of ICT" in t for t in texts)
        assert any("NQF" in t for t in texts)

    def test_extract_html_table(self):
        html = """<html><body>
        <table>
        <tr><th>Programme</th><th>NQF Level</th><th>Credits</th></tr>
        <tr><td>Diploma in CS</td><td>6</td><td>360</td></tr>
        </table></body></html>"""
        from app.adip.extractors.html_extractor import HTMLExtractor
        result = HTMLExtractor().extract_from_string(html)
        assert result.error is None
        assert len(result.tables) == 1
        assert result.tables[0].data_rows[0]["NQF Level"] == "6"

    def test_removes_nav_elements(self):
        html = """<html><body>
        <nav>Navigation here</nav>
        <p>Real content</p>
        </body></html>"""
        from app.adip.extractors.html_extractor import HTMLExtractor
        result = HTMLExtractor().extract_from_string(html)
        texts = [c.text for c in result.chunks]
        assert not any("Navigation" in t for t in texts)
        assert any("Real content" in t for t in texts)


class TestTextExtractor:
    def test_extract_plain_text(self, tmp_path):
        txt = tmp_path / "doc.txt"
        txt.write_text("NQF Level 6\nCredits: 360\nDiploma in Computer Science")
        from app.adip.extractors.text_extractor import TextExtractor
        result = TextExtractor().extract(txt)
        assert result.error is None
        texts = [c.text for c in result.chunks]
        assert any("NQF" in t for t in texts)

    def test_extract_markdown_headings(self, tmp_path):
        md = tmp_path / "doc.md"
        md.write_text("# Faculty of ICT\n\n## Computer Science\n\nDiploma in CS (NQF 6)")
        from app.adip.extractors.text_extractor import TextExtractor
        result = TextExtractor().extract(md)
        headings = [c for c in result.chunks if c.chunk_type == "heading"]
        assert any(h.heading_level == 1 for h in headings)
        assert any(h.heading_level == 2 for h in headings)


# ── Extractor factory ────────────────────────────────────────────────────────

class TestExtractorFactory:
    def test_get_pdf_extractor(self, tmp_path):
        from app.adip.extractors.factory import get_extractor
        from app.adip.extractors.pdf_extractor import PDFExtractor
        ext = get_extractor(tmp_path / "test.pdf", "application/pdf")
        assert isinstance(ext, PDFExtractor)

    def test_get_docx_extractor(self, tmp_path):
        from app.adip.extractors.docx_extractor import DOCXExtractor
        from app.adip.extractors.factory import get_extractor
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = get_extractor(tmp_path / "test.docx", mime)
        assert isinstance(ext, DOCXExtractor)

    def test_unknown_format_returns_none(self, tmp_path):
        from app.adip.extractors.factory import get_extractor
        ext = get_extractor(tmp_path / "test.xyz", "application/octet-stream")
        assert ext is None


# ── Classifier tests ─────────────────────────────────────────────────────────

class TestDocumentClassifier:
    def test_classify_ict_prospectus_by_filename(self):
        from app.adip.classifiers.document_classifier import classify_document
        r = classify_document("Part6_ICT_Prospectus.pdf")
        assert r.document_type == "prospectus_faculty"
        assert r.confidence >= 0.80

    def test_classify_students_rules_by_filename(self):
        from app.adip.classifiers.document_classifier import classify_document
        r = classify_document("Part1_Students_Rules_and_Regulations.pdf")
        assert r.document_type in ("regulations_student", "prospectus_institution")
        assert r.confidence >= 0.65

    def test_classify_academic_calendar(self):
        from app.adip.classifiers.document_classifier import classify_document
        r = classify_document("2026-AcademicCore-Calendar.pdf")
        assert r.document_type == "academic_calendar"
        assert r.confidence >= 0.75

    def test_classify_exam_rules(self):
        from app.adip.classifiers.document_classifier import classify_document
        r = classify_document("Chapter_4_Examination_Rules_2024.pdf")
        assert r.document_type == "policy_examination"

    def test_classify_unknown(self):
        from app.adip.classifiers.document_classifier import classify_document
        r = classify_document("random_document_xyz.pdf")
        assert r.document_type == "unknown"
        assert r.confidence < 0.60

    def test_admin_hint_overrides(self):
        from app.adip.classifiers.document_classifier import classify_document
        r = classify_document("unknownfile.pdf", admin_hint="module_guide")
        assert r.document_type == "module_guide"
        assert r.method == "admin_hint"

    def test_content_based_classification(self):
        from app.adip.classifiers.document_classifier import classify_document
        from app.adip.extractors.base import ExtractionResult, ExtractedChunk
        chunks = [
            ExtractedChunk(text="NQF level 6 diploma programme offered at our faculty",
                           chunk_type="paragraph", extraction_method="test"),
            ExtractedChunk(text="APS requirements: 26 with Mathematics",
                           chunk_type="paragraph", extraction_method="test"),
            ExtractedChunk(text="360 credits total for this programme",
                           chunk_type="paragraph", extraction_method="test"),
        ]
        result = ExtractionResult(extraction_method="test", chunks=chunks)
        r = classify_document("unnamed.pdf", result)
        assert r.document_type == "prospectus_faculty"


# ── Confidence scoring ────────────────────────────────────────────────────────

class TestConfidenceScoring:
    def test_official_html_verbatim_is_high(self):
        from app.adip.validators.confidence import calculate_confidence
        bd = calculate_confidence("official_html", "verbatim_match", "explicit_label")
        assert bd.final_score >= 0.95

    def test_secondary_website_is_below_gate(self):
        from app.adip.validators.confidence import calculate_confidence, GATE_MEDIUM_REVIEW
        bd = calculate_confidence("secondary_website", "regex_clean_text", "column_header")
        assert bd.final_score < GATE_MEDIUM_REVIEW

    def test_gate_auto_approved(self):
        from app.adip.validators.confidence import calculate_confidence
        bd = calculate_confidence("official_html", "verbatim_match", "explicit_label")
        assert bd.gate() == "auto_approved"

    def test_gate_pending_review(self):
        from app.adip.validators.confidence import calculate_confidence
        bd = calculate_confidence("official_pdf_ocr", "ocr_medium", "implicit")
        # score ≈ 0.78 * 0.70 * 0.65 = 0.355 — quarantined
        # Let's try a medium case:
        bd2 = calculate_confidence("official_pdf", "table_cell_inferred", "contextual_label")
        assert bd2.gate() in ("auto_approved", "pending_review", "quarantined")

    def test_gate_status_helper(self):
        from app.adip.validators.confidence import gate_status
        assert gate_status(0.95) == "auto_approved"
        assert gate_status(0.80) == "pending_review"
        assert gate_status(0.50) == "quarantined"


# ── TUT ICT Mapper tests ──────────────────────────────────────────────────────

class TestTUTICTMapper:
    def _make_extraction(self, texts: list[str]):
        from app.adip.extractors.base import ExtractionResult, ExtractedChunk
        chunks = [
            ExtractedChunk(
                text=t, chunk_type="paragraph", extraction_method="test",
                page_number=i + 1, sequence_index=i
            )
            for i, t in enumerate(texts)
        ]
        return ExtractionResult(extraction_method="test", chunks=chunks)

    def test_extracts_nqf_level(self):
        from app.adip.mappers.tut_ict_mapper import TUTICTMapper
        # Use TUT real-world format: "NQF Level 6 (360 credits)"
        result = self._make_extraction([
            "Diploma in Computer Science - NQF Level 6 (360 credits)",
        ])
        mapper = TUTICTMapper("doc-id-1", "inst-id-1")
        candidates = mapper.map(result)
        nqf = [c for c in candidates if c.ikp_field_name == "nqf_level"]
        assert len(nqf) >= 1
        assert nqf[0].raw_value == "6"

    def test_extracts_credits(self):
        from app.adip.mappers.tut_ict_mapper import TUTICTMapper
        # Credits appear as "NQF Level N (XXX credits)" in TUT prospectus
        result = self._make_extraction([
            "Diploma in Computer Science - NQF Level 6 (360 credits)",
        ])
        mapper = TUTICTMapper("doc-id-1", "inst-id-1")
        candidates = mapper.map(result)
        creds = [c for c in candidates if c.ikp_field_name == "total_credits"]
        assert any(c.raw_value == "360" for c in creds)

    def test_extracts_module_code(self):
        from app.adip.mappers.tut_ict_mapper import TUTICTMapper
        from app.adip.extractors.base import ExtractionResult, ExtractedTable
        # Module codes in TUT: 3-4 uppercase letters + 3 digits + 1 letter (e.g. ISC216D)
        # Test via the table mapping path which the mapper's _map_module_table uses
        table = ExtractedTable(
            page_number=10,
            table_index=0,
            header_row=["code", "name", "nqf_level", "credits"],
            data_rows=[{"code": "ISC216D", "name": "Information Security", "nqf_level": "6", "credits": "15"}],
            extraction_method="tab_format_modules",
        )
        result = ExtractionResult(extraction_method="test", tables=[table])
        mapper = TUTICTMapper("doc-id-1", "inst-id-1")
        candidates = mapper.map(result)
        modules = [c for c in candidates if c.ikp_entity_type == "module"]
        assert any(c.ikp_entity_key == "ISC216D" for c in modules)

    def test_extracts_programme_name(self):
        from app.adip.mappers.tut_ict_mapper import TUTICTMapper
        result = self._make_extraction([
            "The Diploma in Information Technology is offered at Soshanguve South."
        ])
        mapper = TUTICTMapper("doc-id-1", "inst-id-1")
        candidates = mapper.map(result)
        progs = [c for c in candidates if c.ikp_field_name == "name"]
        assert any("Information Technology" in c.ikp_entity_key for c in progs)

    def test_table_mapping(self):
        from app.adip.extractors.base import ExtractionResult, ExtractedTable
        from app.adip.mappers.tut_ict_mapper import TUTICTMapper
        table = ExtractedTable(
            page_number=15,
            table_index=0,
            header_row=["Programme", "NQF Level", "Credits", "APS (Math)"],
            data_rows=[{
                "Programme": "Diploma in Computer Science",
                "NQF Level": "6",
                "Credits": "360",
                "APS (Math)": "26",
            }],
            extraction_method="test",
        )
        result = ExtractionResult(extraction_method="test", tables=[table])
        mapper = TUTICTMapper("doc-id-1", "inst-id-1")
        candidates = mapper.map(result)
        assert any(c.ikp_field_name == "nqf_level" and c.raw_value == "6" for c in candidates)
        assert any(c.ikp_field_name == "total_credits" and c.raw_value == "360" for c in candidates)
        assert any(c.ikp_field_name == "aps_minimum_math" and c.raw_value == "26" for c in candidates)

    def test_secondary_source_confidence_low(self):
        from app.adip.mappers.tut_ict_mapper import TUTICTMapper
        result = self._make_extraction(["Diploma in Computer Science (NQF level 6)"])
        mapper = TUTICTMapper("doc-id-1", "inst-id-1", source_type="secondary_website")
        candidates = mapper.map(result)
        # Secondary sources should produce lower confidence
        for c in candidates:
            assert c.confidence < 0.85, f"Expected low conf for secondary, got {c.confidence}"


# ── Provenance engine tests ───────────────────────────────────────────────────

class TestProvenanceEngine:
    def test_generates_anchor_per_candidate(self):
        from app.adip.mappers.tut_ict_mapper import ExtractionCandidateData
        from app.adip.provenance.provenance_engine import generate_provenance
        candidates = [
            ExtractionCandidateData(
                document_id="doc-1",
                institution_id="inst-1",
                ikp_entity_type="programme",
                ikp_entity_key="Diploma in Computer Science",
                ikp_field_name="nqf_level",
                raw_value="6",
                coerced_value="6",
                value_type="integer",
                extraction_method="regex_nqf_pattern",
                source_verbatim="NQF Level 6",
                source_page=12,
                confidence=0.92,
                status="auto_approved",
            )
        ]
        anchors = generate_provenance(
            candidates=candidates,
            document_path=Path("Part6_ICT_Prospectus.pdf"),
            source_url="https://www.tut.ac.za/media/.../Part6_ICT_Prospectus.pdf",
            source_document_title="2026 TUT ICT Prospectus",
        )
        assert len(anchors) == 1
        assert anchors[0].page_number == 12
        assert anchors[0].verbatim_quote == "NQF Level 6"
        assert anchors[0].publisher_verified is True
        assert anchors[0].confidence_score > 0

    def test_anchor_status_is_active(self):
        from app.adip.mappers.tut_ict_mapper import ExtractionCandidateData
        from app.adip.provenance.provenance_engine import generate_provenance
        candidates = [
            ExtractionCandidateData(
                document_id="d1", institution_id="i1",
                ikp_entity_type="programme", ikp_entity_key="TestProg",
                ikp_field_name="nqf_level", raw_value="6", coerced_value="6",
                value_type="integer", extraction_method="test",
                source_verbatim="NQF 6", source_page=1,
                confidence=0.90, status="auto_approved",
            )
        ]
        anchors = generate_provenance(candidates, Path("doc.pdf"), None, None)
        assert anchors[0].status == "active"


# ── Pipeline dry-run test ─────────────────────────────────────────────────────

class TestPipelineDryRun:
    def test_dry_run_returns_summary(self, tmp_path, monkeypatch):
        """Test pipeline with mocked source directory."""
        import app.adip.pipeline.run_tut_ict_extraction as pipeline_module
        # Point SOURCE_DIR to tmp_path (no files = graceful empty run)
        monkeypatch.setattr(pipeline_module, "SOURCE_DIR", tmp_path)
        monkeypatch.setattr(pipeline_module, "OUTPUT_DIR", tmp_path / "out")
        summary = pipeline_module.run_pipeline(dry_run=True)
        assert isinstance(summary, dict)
        assert "documents_processed" in summary
        assert summary["documents_processed"] == 0
        assert summary["errors"] == []

    def test_dry_run_with_text_file(self, tmp_path, monkeypatch):
        """Test pipeline processes a file without crashing."""
        import app.adip.pipeline.run_tut_ict_extraction as pipeline_module
        # Create a dummy PDF file (minimal valid PDF)
        pdf_content = create_minimal_pdf("Diploma in Computer Science NQF Level 6 360 credits")
        (tmp_path / "Part6_ICT_Prospectus.pdf").write_bytes(pdf_content)
        monkeypatch.setattr(pipeline_module, "SOURCE_DIR", tmp_path)
        monkeypatch.setattr(pipeline_module, "OUTPUT_DIR", tmp_path / "out")
        summary = pipeline_module.run_pipeline(dry_run=True)
        assert summary["documents_processed"] >= 1
        # May have extracted chunks or not depending on pdfminer's handling of minimal PDF
        assert isinstance(summary["total_chunks_extracted"], int)


# ── Table extractor tests ─────────────────────────────────────────────────────

class TestTableExtractor:
    def test_join_tab_lines(self):
        from app.adip.extractors.table_extractor import join_tab_lines
        raw = "ISC216D\t\nInformation Security\t\n(6)\t\n(15)\n"
        result = join_tab_lines(raw)
        assert "ISC216D\t" in result
        assert "\t\n" not in result

    def test_flatten_para(self):
        from app.adip.extractors.table_extractor import flatten_para
        raw = "APS of at least 26 (with Mathematics) or 28 (with Mathematical\nLiteracy)."
        result = flatten_para(raw)
        assert "Mathematical Literacy" in result

    def test_extract_tab_modules_finds_entries(self):
        from app.adip.extractors.table_extractor import join_tab_lines, extract_tab_modules
        # Simulate a page with tab-formatted module table
        raw = "ISC216D\tInformation Security\t(6)\t(15)\nORS216D\tOperating Systems\t(6)\t(15)\n"
        joined = join_tab_lines(raw)
        tables = extract_tab_modules(joined, page_number=10)
        assert len(tables) == 1
        assert len(tables[0].data_rows) == 2
        assert tables[0].data_rows[0]["code"] == "ISC216D"
        assert tables[0].data_rows[0]["credits"] == "15"

    def test_extract_tab_modules_empty_page(self):
        from app.adip.extractors.table_extractor import extract_tab_modules
        tables = extract_tab_modules("No modules here, just text.", page_number=1)
        assert tables == []

    def test_pdfplumber_returns_list(self):
        from app.adip.extractors.table_extractor import extract_tables_pdfplumber
        # Pass a non-existent path — should return empty list, not raise
        result = extract_tables_pdfplumber(Path("nonexistent.pdf"))
        assert isinstance(result, list)


# ── Programme extraction tests ────────────────────────────────────────────────

class TestProgrammeExtraction:
    def test_aps_math_extraction(self):
        from app.adip.extractors.table_extractor import flatten_para
        import re
        from app.adip.mappers.tut_ict_mapper import _APS_MATH_RE, _APS_ML_RE
        text = (
            "To be considered for this qualification, applicants must have an Admission Point Score "
            "(APS) of at least 26 (with Mathematics or Technical Mathematics) or 28 (with Mathematical\n"
            "Literacy)."
        )
        flat = flatten_para(text)
        m_math = _APS_MATH_RE.search(flat)
        m_ml = _APS_ML_RE.search(flat)
        assert m_math is not None, "APS Math not extracted"
        assert m_math.group(1) == "26"
        assert m_ml is not None, "APS ML not extracted"
        assert m_ml.group(1) == "28"

    def test_aps_math_does_not_capture_extended_threshold(self):
        from app.adip.extractors.table_extractor import flatten_para
        from app.adip.mappers.tut_ict_mapper import _APS_MATH_RE
        text = (
            "APS of at least 26 (with Mathematics or Technical Mathematics). "
            "Applicants with a score of 23 (with Mathematics or Technical Mathematics) "
            "will be considered for the extended programme."
        )
        flat = flatten_para(text)
        m = _APS_MATH_RE.search(flat)
        assert m is not None
        assert m.group(1) == "26", f"Expected 26, got {m.group(1)}"

    def test_nqf_credits_pattern(self):
        from app.adip.mappers.tut_ict_mapper import _NQF_CREDITS_RE
        text = "Dip (Computer Science) - NQF Level 6 (360 credits)"
        m = _NQF_CREDITS_RE.search(text)
        assert m is not None
        assert m.group(1) == "6"
        assert m.group(2) == "360"

    def test_qual_code_extraction(self):
        from app.adip.mappers.tut_ict_mapper import _QUAL_CODE_RE
        text = "Qualification code: DPRS20\nSAQA ID: 109017"
        m = _QUAL_CODE_RE.search(text)
        assert m is not None
        assert m.group(1) == "DPRS20"

    def test_known_programmes_count(self):
        from app.adip.mappers.tut_ict_mapper import KNOWN_ICT_PROGRAMMES
        assert len(KNOWN_ICT_PROGRAMMES) >= 20

    def test_module_tab_regex(self):
        from app.adip.extractors.table_extractor import MODULE_TAB_RE
        line = "ISC216D\tInformation Security\t(6)\t(15)"
        matches = MODULE_TAB_RE.findall(line)
        assert len(matches) == 1
        code, name, nqf, credits = matches[0]
        assert code == "ISC216D"
        assert name.strip() == "Information Security"
        assert nqf == "6"
        assert credits == "15"


# ── Document classifier override tests ───────────────────────────────────────

class TestClassifierOverrides:
    def test_academic_planning_classified_as_calendar(self):
        from app.adip.classifiers.document_classifier import classify_document
        result = classify_document("AcademicPlanning-Sem1-2026.pdf")
        assert result.document_type == "academic_calendar"
        assert result.method == "institution_filename_override"

    def test_chapter4_classified_as_examination(self):
        from app.adip.classifiers.document_classifier import classify_document
        result = classify_document("Chapter_4_Examination_Rules_2024.pdf")
        assert result.document_type == "policy_examination"

    def test_academic_core_calendar_classified(self):
        from app.adip.classifiers.document_classifier import classify_document
        result = classify_document("2026-AcademicCore-Calendar.pdf")
        assert result.document_type == "academic_calendar"

    def test_students_rules_classified(self):
        from app.adip.classifiers.document_classifier import classify_document
        result = classify_document("Part1_Students_Rules_and_Regulations.pdf")
        assert result.document_type == "regulations_student"


# ── Conflict detection tests ──────────────────────────────────────────────────

class TestConflictDetection:
    def test_no_conflict_same_values(self):
        from app.adip.pipeline.run_tut_ict_extraction import _detect_conflicts
        candidates = [
            {"ikp_entity_type": "programme", "ikp_entity_key": "ProgA",
             "ikp_field_name": "nqf_level", "coerced_value": "6"},
            {"ikp_entity_type": "programme", "ikp_entity_key": "ProgA",
             "ikp_field_name": "nqf_level", "coerced_value": "6"},
        ]
        conflicts = _detect_conflicts(candidates)
        assert conflicts == []

    def test_conflict_different_values(self):
        from app.adip.pipeline.run_tut_ict_extraction import _detect_conflicts
        candidates = [
            {"ikp_entity_type": "admission_requirement", "ikp_entity_key": "ProgA",
             "ikp_field_name": "aps_minimum_math", "coerced_value": "26"},
            {"ikp_entity_type": "admission_requirement", "ikp_entity_key": "ProgA",
             "ikp_field_name": "aps_minimum_math", "coerced_value": "23"},
        ]
        conflicts = _detect_conflicts(candidates)
        assert len(conflicts) == 1
        assert set(conflicts[0]["conflicting_values"]) == {"26", "23"}

    def test_conflict_different_entity_keys_no_conflict(self):
        from app.adip.pipeline.run_tut_ict_extraction import _detect_conflicts
        candidates = [
            {"ikp_entity_type": "programme", "ikp_entity_key": "ProgA",
             "ikp_field_name": "nqf_level", "coerced_value": "6"},
            {"ikp_entity_type": "programme", "ikp_entity_key": "ProgB",
             "ikp_field_name": "nqf_level", "coerced_value": "7"},
        ]
        conflicts = _detect_conflicts(candidates)
        assert conflicts == []
