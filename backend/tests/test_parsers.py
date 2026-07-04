"""Unit tests for document parsers.

Each test uses the respective library to build a minimal valid file in memory
so there are no binary fixtures in the repository.  Tests are skipped
automatically when the required parsing library is not installed.
"""

import asyncio
import csv
import io

import pytest


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def run(coro):
    """Run a coroutine in a new event loop (avoids pytest-asyncio for simplicity)."""
    return asyncio.get_event_loop().run_until_complete(coro)


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------


@pytest.mark.skipif(
    not pytest.importorskip("pypdf", reason="pypdf not installed"),
    reason="pypdf not installed",
)
class TestPdfParser:
    def _make_pdf(self, text: str) -> bytes:
        from pypdf import PdfWriter

        writer = PdfWriter()
        page = writer.add_blank_page(width=595, height=842)
        # Use a minimal content stream to embed text.
        from pypdf.generic import DecodedStreamObject

        content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
        stream = DecodedStreamObject()
        stream.set_data(content)
        page["/Contents"] = stream
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue()

    def test_extracts_page_count(self):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        writer.add_blank_page(width=595, height=842)
        buf = io.BytesIO()
        writer.write(buf)
        content = buf.getvalue()

        from app.parsers.pdf_parser import PdfParser

        result = run(PdfParser().extract(content, "test.pdf"))
        assert result.page_count == 2
        assert result.parser_name == "pdf"

    def test_encrypted_pdf_raises_parser_error_non_recoverable(self):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        writer.encrypt("secret")
        buf = io.BytesIO()
        writer.write(buf)
        content = buf.getvalue()

        from app.parsers.base import ParserError
        from app.parsers.pdf_parser import PdfParser

        with pytest.raises(ParserError) as exc_info:
            run(PdfParser().extract(content, "locked.pdf"))
        assert not exc_info.value.recoverable

    def test_bad_bytes_raises_parser_error(self):
        from app.parsers.base import ParserError
        from app.parsers.pdf_parser import PdfParser

        with pytest.raises(ParserError):
            run(PdfParser().extract(b"not a pdf", "bad.pdf"))


# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------


@pytest.mark.skipif(
    not pytest.importorskip("docx", reason="python-docx not installed"),
    reason="python-docx not installed",
)
class TestDocxParser:
    def _make_docx(self, paragraphs: list[str]) -> bytes:
        import docx

        doc = docx.Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_extracts_paragraphs(self):
        content = self._make_docx(["Hello world", "Second paragraph"])
        from app.parsers.docx_parser import DocxParser

        result = run(DocxParser().extract(content, "test.docx"))
        assert "Hello world" in result.text
        assert "Second paragraph" in result.text
        assert result.word_count >= 4

    def test_extracts_table_data(self):
        import docx

        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Score"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "85"
        buf = io.BytesIO()
        doc.save(buf)

        from app.parsers.docx_parser import DocxParser

        result = run(DocxParser().extract(buf.getvalue(), "marks.docx"))
        assert "Alice" in result.text
        assert result.metadata["table_count"] == 1

    def test_empty_document(self):
        import docx

        doc = docx.Document()
        buf = io.BytesIO()
        doc.save(buf)

        from app.parsers.docx_parser import DocxParser

        result = run(DocxParser().extract(buf.getvalue(), "empty.docx"))
        assert result.word_count == 0


# ---------------------------------------------------------------------------
# XLSX parser
# ---------------------------------------------------------------------------


@pytest.mark.skipif(
    not pytest.importorskip("openpyxl", reason="openpyxl not installed"),
    reason="openpyxl not installed",
)
class TestXlsxParser:
    def _make_xlsx(self, rows: list[list]) -> bytes:
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        for row in rows:
            ws.append(row)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_extracts_cell_text(self):
        content = self._make_xlsx([["Name", "Mark"], ["Bob", 72], ["Carol", 88]])
        from app.parsers.xlsx_parser import XlsxParser

        result = run(XlsxParser().extract(content, "marks.xlsx"))
        assert "Bob" in result.text
        assert "Carol" in result.text
        assert result.metadata["sheet_count"] == 1
        assert result.metadata["first_sheet_headers"] == ["Name", "Mark"]

    def test_multiple_sheets(self):
        import openpyxl

        wb = openpyxl.Workbook()
        wb.active.title = "Sheet1"
        wb.create_sheet("Sheet2")
        buf = io.BytesIO()
        wb.save(buf)

        from app.parsers.xlsx_parser import XlsxParser

        result = run(XlsxParser().extract(buf.getvalue(), "multi.xlsx"))
        assert result.page_count == 2
        assert "Sheet1" in result.metadata["sheet_names"]
        assert "Sheet2" in result.metadata["sheet_names"]


# ---------------------------------------------------------------------------
# TXT parser
# ---------------------------------------------------------------------------


class TestTxtParser:
    def test_plain_text(self):
        from app.parsers.text_parser import TxtParser

        content = "Hello world\nThis is line two\n".encode("utf-8")
        result = run(TxtParser().extract(content, "doc.txt"))
        assert "Hello world" in result.text
        assert result.word_count >= 6
        assert result.metadata["line_count"] == 2

    def test_utf8_bom(self):
        from app.parsers.text_parser import TxtParser

        content = "﻿Hello BOM".encode("utf-8-sig")
        result = run(TxtParser().extract(content, "bom.txt"))
        assert "Hello BOM" in result.text

    def test_empty_file_returns_zero_word_count(self):
        from app.parsers.text_parser import TxtParser

        result = run(TxtParser().extract(b"", "empty.txt"))
        assert result.word_count == 0


# ---------------------------------------------------------------------------
# CSV parser
# ---------------------------------------------------------------------------


class TestCsvParser:
    def _make_csv(self, rows: list[list[str]]) -> bytes:
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerows(rows)
        return buf.getvalue().encode("utf-8")

    def test_extracts_rows(self):
        from app.parsers.text_parser import CsvParser

        content = self._make_csv([["Name", "Grade"], ["Alice", "A"], ["Bob", "B"]])
        result = run(CsvParser().extract(content, "grades.csv"))
        assert "Alice" in result.text
        assert result.metadata["row_count"] == 3
        assert result.metadata["headers"] == ["Name", "Grade"]

    def test_empty_csv(self):
        from app.parsers.text_parser import CsvParser

        result = run(CsvParser().extract(b"", "empty.csv"))
        assert result.word_count == 0


# ---------------------------------------------------------------------------
# Image parser (OCR graceful degradation)
# ---------------------------------------------------------------------------


class TestImageParser:
    def test_returns_result_when_ocr_unavailable(self):
        """ImageParser must not raise even when pytesseract/Pillow are absent."""
        from app.parsers.image_parser import ImageParser

        # Use a 1×1 white PNG (minimal valid PNG bytes).
        minimal_png = (
            b"\x89PNG\r\n\x1a\n"
            b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDAT"
            b"x\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18\xd8N"
            b"\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        result = run(ImageParser().extract(minimal_png, "test.png"))
        assert result.parser_name == "image"
        assert isinstance(result.text, str)


# ---------------------------------------------------------------------------
# Parser factory
# ---------------------------------------------------------------------------


class TestParserFactory:
    def test_returns_correct_parser_for_pdf(self):
        from app.parsers.factory import get_parser
        from app.parsers.pdf_parser import PdfParser

        assert isinstance(get_parser("application/pdf"), PdfParser)

    def test_raises_key_error_for_unknown_mime(self):
        from app.parsers.factory import get_parser

        with pytest.raises(KeyError):
            get_parser("application/x-unknown-format")

    def test_is_supported_returns_false_for_unknown(self):
        from app.parsers.factory import is_supported

        assert not is_supported("application/x-unknown-format")

    def test_all_mime_types_covered(self):
        from app.parsers.factory import supported_mime_types

        types = supported_mime_types()
        assert "application/pdf" in types
        assert "image/png" in types
        assert "text/csv" in types
        assert "application/zip" in types
