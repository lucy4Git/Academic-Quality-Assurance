"""DOCX parser using python-docx.

Extracts body paragraphs, table cell text, and heading-style paragraphs.
Heading paragraphs (style names starting with "Heading") are also surfaced
separately in metadata so the classification service can weight them more
heavily.
"""

from __future__ import annotations

import io

from app.parsers.base import DocumentParser, ExtractionResult, ParserError


class DocxParser(DocumentParser):
    """Extract text and tables from DOCX files."""

    @property
    def supported_mime_types(self) -> frozenset[str]:
        return frozenset(
            {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        )

    async def extract(self, content: bytes, filename: str) -> ExtractionResult:
        return await self._run_sync(self._extract_sync, content, filename)

    @staticmethod
    def _extract_sync(content: bytes, filename: str) -> ExtractionResult:
        try:
            import docx  # noqa: PLC0415
        except ImportError:
            raise ParserError(
                "python-docx is not installed. Run: pip install python-docx",
                recoverable=False,
            )

        try:
            doc = docx.Document(io.BytesIO(content))
        except Exception as exc:
            raise ParserError(f"Cannot open DOCX '{filename}': {exc}") from exc

        parts: list[str] = []
        headings: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            parts.append(text)
            if para.style and para.style.name.startswith("Heading"):
                headings.append(text)

        # Extract table cells
        tables_data: list[list[list[str]]] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
                parts.extend(c for c in cells if c)
            tables_data.append(rows)

        full_text = "\n".join(parts).strip()

        return ExtractionResult(
            text=full_text,
            word_count=len(full_text.split()),
            parser_name="docx",
            page_count=None,  # python-docx does not expose page count
            metadata={
                "headings": headings,
                "table_count": len(tables_data),
                "tables": tables_data[:10],  # store first 10 tables max
            },
        )
