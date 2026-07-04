"""ADIP DOCX extractor using python-docx."""

from __future__ import annotations

from pathlib import Path

from app.adip.extractors.base import (
    BaseExtractor,
    DocumentMetadata,
    ExtractedChunk,
    ExtractedTable,
    ExtractionResult,
)

SUPPORTED_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
METHOD = "python_docx"


class DOCXExtractor(BaseExtractor):
    """Extracts paragraphs, headings, and tables from DOCX files."""

    def can_handle(self, mime_type: str, file_path: Path) -> bool:
        return mime_type in SUPPORTED_MIMES or str(file_path).lower().endswith((".docx", ".doc"))

    def extract(self, file_path: Path) -> ExtractionResult:
        try:
            import docx
        except ImportError:
            return ExtractionResult(
                extraction_method=METHOD,
                error="python-docx not installed. Run: pip install python-docx",
            )

        chunks: list[ExtractedChunk] = []
        tables: list[ExtractedTable] = []
        section_path: list[str] = []
        sequence = 0

        try:
            doc = docx.Document(str(file_path))
        except Exception as exc:
            return ExtractionResult(extraction_method=METHOD, error=f"Cannot open DOCX: {exc}")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            heading_level: int | None = None
            chunk_type = "paragraph"

            if style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    heading_level = 1
                chunk_type = "heading"
                depth = heading_level - 1
                section_path = section_path[:depth]
                section_path.append(text)

            chunks.append(ExtractedChunk(
                text=text,
                chunk_type=chunk_type,
                section_path=list(section_path),
                heading_level=heading_level,
                extraction_method=METHOD,
                sequence_index=sequence,
            ))
            sequence += 1

        for t_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            data_rows = []
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    data_rows.append({headers[i]: cells[i] if i < len(cells) else "" for i in range(len(headers))})
            tables.append(ExtractedTable(
                page_number=None,
                table_index=t_idx,
                header_row=headers,
                data_rows=data_rows,
                extraction_method=METHOD,
            ))

        props = doc.core_properties
        metadata = DocumentMetadata(
            title=props.title or None,
            author=props.author or None,
        )

        return ExtractionResult(
            extraction_method=METHOD,
            chunks=chunks,
            tables=tables,
            metadata=metadata,
            extraction_quality=0.93,
        )
